#!/usr/bin/env python3
"""
Generates the Ubumenyi Technical Architecture Document (.docx)
suitable for REB proposal appendices and institutional presentations.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Color constants ──────────────────────────────────────────
BLUE       = RGBColor(0x34, 0x57, 0xA6)   # Ubumenyi brand blue
DARK       = RGBColor(0x0F, 0x17, 0x2A)   # Near-black
MID        = RGBColor(0x33, 0x41, 0x55)
MUTED      = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BG   = RGBColor(0xF8, 0xFA, 0xFC)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BORDER     = RGBColor(0xE2, 0xE8, 0xF0)
GREEN      = RGBColor(0x05, 0x96, 0x69)
AMBER      = RGBColor(0x92, 0x40, 0x0E)


def hex_color(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


BLUE_HEX   = hex_color(0x34, 0x57, 0xA6)
DARK_HEX   = hex_color(0x0F, 0x17, 0x2A)
MID_HEX    = hex_color(0x33, 0x41, 0x55)
MUTED_HEX  = hex_color(0x64, 0x74, 0x8B)
LBG_HEX    = hex_color(0xF8, 0xFA, 0xFC)
GREEN_HEX  = hex_color(0x05, 0x96, 0x69)
AMBER_HEX  = hex_color(0xD9, 0x77, 0x06)
RED_HEX    = hex_color(0xDC, 0x26, 0x26)


def set_cell_background(cell, hex_color_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color_str)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', 'E2E8F0'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK if level > 1 else BLUE
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(20)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.bold = True
    return p


def add_body(doc, text, color=None, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color or MID
    run.italic = italic
    return p


def add_bullet(doc, text, indent_level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + indent_level * 0.25)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10.5)
        run_b.font.color.rgb = DARK
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = MID
    return p


def add_table_header_row(table, headers, bg_hex=BLUE_HEX, text_color=WHITE):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_background(cell, bg_hex)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = text_color
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)


def add_table_data_row(table, values, row_idx, alt=False):
    row = table.add_row()
    bg = LBG_HEX if alt else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        set_cell_background(cell, bg)
        p = cell.paragraphs[0]
        if isinstance(val, tuple):
            text, bold = val
        else:
            text, bold = val, False
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK if bold else MID
        run.font.bold = bold
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)


def add_section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DBEAFE')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_info_box(doc, title, lines, bg_hex='EFF6FF', border_hex='BFDBFE', title_color=None):
    """Add a styled info box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_border(cell,
        top={'val': 'single', 'sz': 6, 'color': border_hex},
        bottom={'val': 'single', 'sz': 6, 'color': border_hex},
        left={'val': 'single', 'sz': 6, 'color': border_hex},
        right={'val': 'single', 'sz': 6, 'color': border_hex},
    )
    # Title
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(title)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = title_color or BLUE
    p.paragraph_format.space_before = Pt(4)
    # Lines
    for line in lines:
        lp = cell.add_paragraph()
        r = lp.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = MID
    cell.paragraphs[-1].paragraph_format.space_after = Pt(4)
    doc.add_paragraph()  # spacing after


# ═══════════════════════════════════════════════════════════════
# DOCUMENT START
# ═══════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# ── Cover / Title page ─────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('UBUMENYI')
r.font.size = Pt(32)
r.font.bold = True
r.font.name = 'Calibri'
r.font.color.rgb = BLUE

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Educational Resource Initiative')
r2.font.size = Pt(16)
r2.font.name = 'Calibri'
r2.font.color.rgb = MID

doc.add_paragraph()

p_doc = doc.add_paragraph()
p_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_doc.add_run('Technical Architecture & Platform Documentation')
r3.font.size = Pt(13)
r3.font.bold = True
r3.font.name = 'Calibri'
r3.font.color.rgb = DARK

doc.add_paragraph()
doc.add_paragraph()

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Version 2.0   |   May 2025',
    'Kigali, Rwanda',
    'Prepared for: Rwanda Education Board (REB) Proposal Appendix',
]:
    rr = p_meta.add_run(line + '\n')
    rr.font.size = Pt(10)
    rr.font.name = 'Calibri'
    rr.font.color.rgb = MUTED

doc.add_page_break()


# ── 1. Executive Summary ───────────────────────────────────────
add_heading(doc, '1. Executive Summary', level=1)
add_section_divider(doc)

add_body(doc,
    'Ubumenyi is Rwanda\'s first centralized educational resource infrastructure — a digital platform '
    'designed to organize, verify, and freely distribute educational materials to every student, '
    'teacher, and institution in the country. The platform operates as a trusted public educational '
    'utility: not a commercial product, not an open-upload website, but a carefully moderated '
    'institutional resource system aligned to the Rwanda Education Board (REB) curriculum.',
    size=10.5
)

doc.add_paragraph()
add_body(doc,
    'This document provides a complete technical and architectural overview of the Ubumenyi platform '
    'for institutional review, including the database schema, contributor verification system, '
    'moderation workflows, technology stack, and design principles.',
    size=10.5
)

doc.add_paragraph()
add_info_box(doc, 'Core Platform Principles', [
    '• Every resource is reviewed and verified before publication — no unmoderated uploads',
    '• Every contributor is individually screened and approved before publishing',
    '• All resources are freely accessible — no paywalls, no registration required to browse',
    '• Platform is mobile-first and optimized for low-bandwidth environments',
    '• Architecture is built for institutional scale (10,000+ resources)',
    '• Curriculum-aligned: all content organized by REB level, subject, and topic taxonomy',
])


# ── 2. Technology Stack ────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '2. Technology Stack', level=1)
add_section_divider(doc)

add_body(doc,
    'Ubumenyi is built on a modern, production-grade open-source stack selected for reliability, '
    'developer accessibility, scalability, and low operational cost.',
    size=10.5
)
doc.add_paragraph()

# Stack table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.autofit = False
table.columns[0].width = Inches(1.6)
table.columns[1].width = Inches(2.0)
table.columns[2].width = Inches(2.8)
add_table_header_row(table, ['Layer', 'Technology', 'Purpose'])

stack_rows = [
    ('Frontend', 'Next.js 15 (React 19)', 'Server-side rendering, routing, API routes — optimized for performance on slow connections'),
    ('Styling', 'Tailwind CSS v4', 'Utility-first CSS — mobile-first responsive design with minimal bundle size'),
    ('UI Components', 'shadcn/ui + Radix UI', 'Accessible, unstyled component primitives — composable and lightweight'),
    ('Backend / DB', 'Supabase (PostgreSQL)', 'Managed Postgres database with real-time, row-level security, and built-in auth'),
    ('Storage', 'Supabase Storage', 'File hosting for PDFs and documents with role-based access policies'),
    ('Auth', 'Supabase Auth', 'Email authentication with admin/contributor role management via RLS'),
    ('Language', 'TypeScript 5', 'Full type safety across frontend and server actions'),
    ('Hosting', 'Vercel (Edge)', 'Global CDN deployment — fast page loads in Rwanda and across Africa'),
    ('Package Manager', 'npm / Node.js 20+', 'Standard Node.js toolchain'),
]

for i, (layer, tech, purpose) in enumerate(stack_rows):
    add_table_data_row(table, [(layer, True), (tech, False), (purpose, False)], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 3. Database Architecture ───────────────────────────────────
add_heading(doc, '3. Database Architecture', level=1)
add_section_divider(doc)

add_body(doc,
    'The database is built on PostgreSQL via Supabase, with Row-Level Security (RLS) enforced at the '
    'database layer for every table. This ensures that authorization is database-enforced, not just '
    'application-level — a critical requirement for a platform handling publicly accessible and '
    'institution-attributed educational content.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '3.1 Core Tables Overview', level=2)
doc.add_paragraph()

# Tables overview table
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
table2.autofit = False
table2.columns[0].width = Inches(1.8)
table2.columns[1].width = Inches(1.6)
table2.columns[2].width = Inches(3.0)
add_table_header_row(table2, ['Table', 'Type', 'Description'])

db_tables = [
    ('resources', 'Core', 'Educational materials — past papers, syllabi, marking schemes, books, etc.'),
    ('profiles', 'Auth/Users', 'User accounts with role (admin/contributor/reviewer) and verification flags'),
    ('contributor_profiles', 'Contributors', 'Verified contributor profiles: bio, institution, specialties, verification status'),
    ('contributor_applications', 'Workflow', 'Applications from users seeking contributor status — pending review queue'),
    ('topics', 'Taxonomy', 'Reference table of subjects and topics for resource classification'),
    ('resource_downloads', 'Analytics', 'Download event log for engagement tracking and analytics'),
    ('papers', 'Legacy', 'Interactive past paper questions for the Umukoro assessment platform'),
    ('questions', 'Legacy', 'Individual past paper questions with MCQ/short/long types'),
    ('attempts', 'Legacy', 'Student practice attempt records linked to papers'),
    ('notes', 'Legacy', 'Contributor-uploaded study notes linked to subjects'),
    ('syllabi', 'Reference', 'Official curriculum syllabi content linked to subjects'),
    ('subjects', 'Reference', 'Subject taxonomy with level, name, and color coding'),
    ('saves', 'Engagement', 'User-saved resource bookmarks'),
    ('likes', 'Engagement', 'Resource likes/upvotes'),
]

for i, (tbl, typ, desc) in enumerate(db_tables):
    add_table_data_row(table2, [(tbl, True), (typ, False), (desc, False)], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '3.2 Resources Table — Full Schema', level=2)
add_body(doc,
    'The resources table is the core of the platform. It stores all educational materials with '
    'rich metadata for search, filtering, and moderation. All fields are indexed appropriately for '
    'the filtering patterns described in Section 5.',
    size=10.5
)
doc.add_paragraph()

# Resources schema table
table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
table3.autofit = False
table3.columns[0].width = Inches(1.6)
table3.columns[1].width = Inches(1.3)
table3.columns[2].width = Inches(0.8)
table3.columns[3].width = Inches(2.7)
add_table_header_row(table3, ['Column', 'Type', 'Required', 'Description'])

resources_schema = [
    ('id', 'UUID', 'Yes', 'Auto-generated primary key'),
    ('title', 'TEXT', 'Yes', 'Full title of the resource'),
    ('description', 'TEXT', 'No', 'Optional summary/notes about the resource'),
    ('resource_type', 'TEXT', 'Yes', 'past_paper | marking_scheme | syllabus | teacher_notes | book | curriculum | educational_guide'),
    ('level', 'TEXT', 'Yes', 'primary | o_level | a_level | tvet'),
    ('subject', 'TEXT', 'No', 'Subject (e.g. Mathematics, Chemistry)'),
    ('topic', 'TEXT', 'No', 'Specific topic or chapter within the subject'),
    ('combination', 'TEXT', 'No', 'A-Level subject combination (e.g. MCB, MPC) — A-Level only'),
    ('language', 'TEXT', 'No', 'english | french | kinyarwanda | bilingual — default: english'),
    ('tags', 'TEXT[]', 'No', 'Free-form tag array for additional classification (max 10)'),
    ('year', 'INTEGER', 'No', 'Academic year (1990–2099), e.g. 2023'),
    ('is_official', 'BOOLEAN', 'No', 'True if this is an official REB/government document'),
    ('source_institution', 'TEXT', 'No', 'Name of the official source (if is_official = true)'),
    ('file_url', 'TEXT', 'No', 'Public URL to the file in Supabase Storage'),
    ('file_name', 'TEXT', 'No', 'Original filename (e.g. math_paper1_2023.pdf)'),
    ('file_size', 'BIGINT', 'No', 'File size in bytes'),
    ('status', 'TEXT', 'Yes', 'draft | pending_review | approved | rejected | published | archived'),
    ('uploaded_by', 'UUID → profiles', 'No', 'FK to profiles.id — the user who uploaded'),
    ('reviewed_by', 'UUID → profiles', 'No', 'FK to profiles.id — first reviewer'),
    ('moderator_id', 'UUID → profiles', 'No', 'FK to profiles.id — moderator who made final decision'),
    ('rejection_reason', 'TEXT', 'No', 'Reason provided when status = rejected'),
    ('moderation_notes', 'TEXT', 'No', 'Internal notes from the moderation team'),
    ('moderation_at', 'TIMESTAMPTZ', 'No', 'Timestamp of moderation decision'),
    ('download_count', 'INTEGER', 'No', 'Total downloads — updated by database function, default 0'),
    ('view_count', 'INTEGER', 'No', 'Total page views — updated by database function, default 0'),
    ('published_at', 'TIMESTAMPTZ', 'No', 'Timestamp when status changed to published'),
    ('created_at', 'TIMESTAMPTZ', 'No', 'Record creation timestamp — auto-set'),
    ('updated_at', 'TIMESTAMPTZ', 'No', 'Last update timestamp — auto-maintained by trigger'),
]

for i, row in enumerate(resources_schema):
    add_table_data_row(table3, [(row[0], True), row[1], row[2], row[3]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '3.3 Contributor Profiles Table', level=2)
add_body(doc,
    'Every verified contributor has a contributor_profile record separate from their auth profile. '
    'This supports multiple contributor types (individual teachers, schools, institutions, educational '
    'organizations) and stores all public-facing information.',
    size=10.5
)
doc.add_paragraph()

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
table4.autofit = False
table4.columns[0].width = Inches(1.9)
table4.columns[1].width = Inches(1.5)
table4.columns[2].width = Inches(3.0)
add_table_header_row(table4, ['Column', 'Type', 'Description'])

contributor_schema = [
    ('id', 'UUID', 'Primary key'),
    ('user_id', 'UUID → profiles', 'Unique FK to auth profiles — one profile per user'),
    ('display_name', 'TEXT', 'Public-facing name (person, school, or organization name)'),
    ('bio', 'TEXT', 'Optional biography or description'),
    ('avatar_url', 'TEXT', 'Profile photo URL'),
    ('institution_name', 'TEXT', 'School, university, or organization name'),
    ('institution_type', 'TEXT', 'school / university / ministry / NGO / individual'),
    ('location', 'TEXT', 'Location string, e.g. "Kigali, Rwanda"'),
    ('website_url', 'TEXT', 'Optional website URL'),
    ('subject_specialties', 'TEXT[]', 'Array of subjects the contributor specializes in'),
    ('contributor_type', 'TEXT', 'teacher | school | institution | educational_org | individual'),
    ('verification_status', 'TEXT', 'pending | verified | rejected | suspended'),
    ('verified_at', 'TIMESTAMPTZ', 'Timestamp of verification'),
    ('verified_by', 'UUID → profiles', 'Admin who performed the verification'),
    ('verification_notes', 'TEXT', 'Admin notes from the verification process'),
    ('published_count', 'INTEGER', 'Auto-maintained count of published resources — updated by trigger'),
    ('created_at', 'TIMESTAMPTZ', 'Record creation timestamp'),
    ('updated_at', 'TIMESTAMPTZ', 'Auto-maintained last update timestamp'),
]

for i, row in enumerate(contributor_schema):
    add_table_data_row(table4, [(row[0], True), row[1], row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '3.4 Database Security (Row-Level Security)', level=2)
add_body(doc,
    'Every table has RLS enabled. Policies are defined per operation (SELECT, INSERT, UPDATE, DELETE) '
    'and role. Key security rules:',
    size=10.5
)
doc.add_paragraph()

rls_rules = [
    ('Resources', 'SELECT (public)', 'Anyone can view resources with status = published'),
    ('Resources', 'SELECT (auth)', 'Authenticated users (admins) can view all resources'),
    ('Resources', 'INSERT', 'Authenticated users can insert — only setting uploaded_by = auth.uid()'),
    ('Resources', 'UPDATE', 'Only the uploader or an admin can update a resource'),
    ('Resources', 'DELETE', 'Only admins can delete resources'),
    ('Contributor Profiles', 'SELECT (public)', 'Anyone can view profiles with verification_status = verified'),
    ('Contributor Profiles', 'INSERT', 'Users can create their own contributor profile only'),
    ('Contributor Profiles', 'UPDATE', 'Users can update their own profile; admins can update any'),
    ('Topics', 'SELECT', 'Public read access'),
    ('Topics', 'ALL', 'Admin-only write access'),
    ('Resource Downloads', 'INSERT', 'Any authenticated user can log a download event'),
    ('Resource Downloads', 'SELECT', 'Only admins can read the download log'),
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
table5.autofit = False
table5.columns[0].width = Inches(1.7)
table5.columns[1].width = Inches(1.5)
table5.columns[2].width = Inches(3.2)
add_table_header_row(table5, ['Table', 'Operation', 'Policy'])

for i, row in enumerate(rls_rules):
    add_table_data_row(table5, [(row[0], True), row[1], row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 4. Contributor System ──────────────────────────────────────
doc.add_page_break()
add_heading(doc, '4. Contributor Verification System', level=1)
add_section_divider(doc)

add_body(doc,
    'The Ubumenyi contributor system is not an open-upload platform. Every contributor must apply, '
    'be reviewed, and be approved by the moderation team before they can publish any resources. This '
    'ensures that the platform maintains institutional-grade content quality.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '4.1 Contributor Types', level=2)

contributor_types = [
    ('Teacher', 'Individual certified educators with subject expertise'),
    ('School', 'Recognized Rwandan schools contributing institutional materials'),
    ('Institution', 'Universities, colleges, and educational institutions'),
    ('Educational Organization', 'NGOs and educational organizations with curriculum expertise'),
    ('Individual Contributor', 'Verified individuals with demonstrated academic credentials'),
]

table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Table Grid'
table6.autofit = False
table6.columns[0].width = Inches(2.2)
table6.columns[1].width = Inches(4.2)
add_table_header_row(table6, ['Type', 'Description'])

for i, row in enumerate(contributor_types):
    add_table_data_row(table6, [(row[0], True), row[1]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '4.2 Verification Workflow', level=2)
doc.add_paragraph()

steps = [
    ('Step 1', 'Application Submission', 'The applicant completes a contributor application including: motivation, school/organization affiliation, and subject expertise areas.'),
    ('Step 2', 'Admin Review Queue', 'The application appears in the Admin → Contributors panel, flagged as pending. Admins receive a notification banner.'),
    ('Step 3', 'Identity & Credential Check', 'The moderation team reviews the applicant\'s stated credentials, institution, and motivation against available records.'),
    ('Step 4', 'Approve or Reject', 'Admin approves (sets verification_status = verified, updates profiles.is_contributor = true) or rejects with optional notes.'),
    ('Step 5', 'Profile Published', 'Approved contributors appear on the public Contributors directory and can begin uploading resources.'),
    ('Step 6', 'Upload Moderation', 'Every uploaded resource still requires moderation approval before publication — contributor status does not bypass content review.'),
]

table7 = doc.add_table(rows=1, cols=3)
table7.style = 'Table Grid'
table7.autofit = False
table7.columns[0].width = Inches(0.7)
table7.columns[1].width = Inches(1.8)
table7.columns[2].width = Inches(3.9)
add_table_header_row(table7, ['Step', 'Stage', 'Description'])

for i, row in enumerate(steps):
    add_table_data_row(table7, [(row[0], True), (row[1], True), row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 5. Resource Upload & Moderation Workflow ───────────────────
add_heading(doc, '5. Resource Upload & Moderation Workflow', level=1)
add_section_divider(doc)

add_body(doc,
    'Resources undergo a multi-stage workflow from upload to publication. The workflow is designed to '
    'maintain quality, prevent low-quality or inappropriate content, and ensure accurate metadata '
    'for searchability.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '5.1 Resource Status Lifecycle', level=2)
doc.add_paragraph()

statuses = [
    ('draft', 'Saved by contributor — not yet submitted for review. Not visible publicly.'),
    ('pending_review', 'Submitted for review. Visible to admins in the moderation queue.'),
    ('approved', 'Reviewed and approved but not yet set to published. Optional staging state.'),
    ('rejected', 'Rejected with a reason. Uploader is notified. Not visible publicly.'),
    ('published', 'Fully reviewed and live. Visible to all users in the resource library.'),
    ('archived', 'Previously published but removed from active listing. Not visible publicly.'),
]

table8 = doc.add_table(rows=1, cols=2)
table8.style = 'Table Grid'
table8.autofit = False
table8.columns[0].width = Inches(1.5)
table8.columns[1].width = Inches(4.9)
add_table_header_row(table8, ['Status', 'Description'])

for i, row in enumerate(statuses):
    add_table_data_row(table8, [(row[0], True), row[1]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '5.2 Moderation Capabilities', level=2)
add_body(doc, 'Moderators have access to the following actions in the Admin panel:', size=10.5)
doc.add_paragraph()

moderation_actions = [
    'Preview PDF resources directly in the admin panel before approving',
    'Edit all metadata fields (title, subject, topic, combination, year, tags, language)',
    'Publish a resource (sets status = published with timestamp and moderator_id)',
    'Reject a resource with a written reason (stored in rejection_reason column)',
    'Add internal moderation notes (visible only to admins — moderation_notes column)',
    'Archive a published resource (removes from public listing without deletion)',
    'Permanently delete a resource (also removes the file from Supabase Storage)',
    'Filter resources by status, type, subject, or search by title',
]

for action in moderation_actions:
    add_bullet(doc, action)

doc.add_paragraph()

add_heading(doc, '5.3 Resource Classification Schema', level=2)
add_body(doc,
    'Every resource is classified across multiple dimensions to support precise search and filtering:',
    size=10.5
)
doc.add_paragraph()

classifications = [
    ('Level', 'primary | o_level | a_level | tvet', 'Top-level education stage'),
    ('Subject', 'Mathematics, Physics, Chemistry, etc.', 'Subject within the level'),
    ('Topic', 'Free text', 'Specific chapter or topic within the subject'),
    ('Combination', 'MCB, MPC, HEG, etc.', 'A-Level subject combination (A-Level only)'),
    ('Resource Type', '7 types', 'Past paper, syllabus, marking scheme, etc.'),
    ('Year', '1990–2099', 'Academic year the resource applies to'),
    ('Language', 'english | french | kinyarwanda | bilingual', 'Language of the resource'),
    ('Tags', 'Array of up to 10 strings', 'Free-form additional classification tags'),
    ('is_official', 'Boolean', 'Whether this is an official REB/government document'),
]

table9 = doc.add_table(rows=1, cols=3)
table9.style = 'Table Grid'
table9.autofit = False
table9.columns[0].width = Inches(1.4)
table9.columns[1].width = Inches(2.4)
table9.columns[2].width = Inches(2.6)
add_table_header_row(table9, ['Dimension', 'Values', 'Notes'])

for i, row in enumerate(classifications):
    add_table_data_row(table9, [(row[0], True), row[1], row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 6. Platform Pages & Features ──────────────────────────────
doc.add_page_break()
add_heading(doc, '6. Platform Pages & Features', level=1)
add_section_divider(doc)

add_heading(doc, '6.1 Public-Facing Pages', level=2)
doc.add_paragraph()

pages = [
    ('/', 'Landing Page', 'Mission statement, resource type overview, contributor CTA, platform description, partner contact section'),
    ('/resources', 'Resource Library', 'Full searchable and filterable catalog of published resources with sidebar filters (level, type, subject, topic, combination, year, language, official only) and active filter chips'),
    ('/resources/[id]', 'Resource Detail', 'Full resource metadata, embedded PDF viewer, download button, contributor attribution card, related resources panel, view/download counts, tag navigation'),
    ('/contributors', 'Contributors Directory', 'Public directory of verified contributors with search and filter by type and subject'),
    ('/contributors/[id]', 'Contributor Profile', 'Public profile page: bio, institution, verification badge, subject specialties, published resources listing'),
]

table10 = doc.add_table(rows=1, cols=3)
table10.style = 'Table Grid'
table10.autofit = False
table10.columns[0].width = Inches(1.6)
table10.columns[1].width = Inches(1.5)
table10.columns[2].width = Inches(3.3)
add_table_header_row(table10, ['Route', 'Page', 'Key Features'])

for i, row in enumerate(pages):
    add_table_data_row(table10, [(row[0], True), (row[1], True), row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '6.2 Admin Panel Pages', level=2)
doc.add_paragraph()

admin_pages = [
    ('/admin', 'Dashboard', 'Resource counts by status, type breakdown chart, recent uploads, pending review alert'),
    ('/admin/resources', 'Resource Management', 'Full resource table with status tabs, search, publish/reject/archive/delete actions, rejection reason prompt'),
    ('/admin/resources/new', 'Upload Resource', 'Full-featured upload form: all classification fields, file upload with progress, draft/submit-for-review options'),
    ('/admin/resources/[id]/edit', 'Edit Resource', 'Identical to upload form but pre-populated — replaces file optionally'),
    ('/admin/contributors', 'Contributors', 'Contributor directory with verification status tabs, pending applications panel, approve/reject actions'),
    ('/admin/team', 'Team', 'Admin user management — invite new admins'),
]

table11 = doc.add_table(rows=1, cols=3)
table11.style = 'Table Grid'
table11.autofit = False
table11.columns[0].width = Inches(1.9)
table11.columns[1].width = Inches(1.5)
table11.columns[2].width = Inches(3.0)
add_table_header_row(table11, ['Route', 'Page', 'Key Features'])

for i, row in enumerate(admin_pages):
    add_table_data_row(table11, [(row[0], True), (row[1], True), row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 7. Search & Filter Architecture ────────────────────────────
add_heading(doc, '7. Search & Filter Architecture', level=1)
add_section_divider(doc)

add_body(doc,
    'The resource search and filter system is implemented as server-side URL-driven filtering — '
    'meaning all filter state is stored in the URL query string, making results shareable and '
    'bookmarkable. No client-side JavaScript filtering is used, which ensures performance on '
    'low-powered devices and slow networks.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '7.1 Available Filter Dimensions', level=2)
doc.add_paragraph()

filters = [
    ('q', 'Text search', 'ilike pattern match on resource title — case-insensitive'),
    ('level', 'Education level', 'Exact match on level column'),
    ('type', 'Resource type', 'Exact match on resource_type column'),
    ('subject', 'Subject', 'Exact match on subject column'),
    ('topic', 'Topic', 'ilike pattern match on topic column'),
    ('combination', 'A-Level combination', 'Exact match — shown only when level=a_level'),
    ('year', 'Academic year', 'Exact integer match on year column'),
    ('language', 'Language', 'Exact match on language column'),
    ('official', 'Official documents', 'Filters to is_official = true when set to "1"'),
]

table12 = doc.add_table(rows=1, cols=3)
table12.style = 'Table Grid'
table12.autofit = False
table12.columns[0].width = Inches(1.2)
table12.columns[1].width = Inches(1.6)
table12.columns[2].width = Inches(3.6)
add_table_header_row(table12, ['URL Param', 'Filter', 'DB Operation'])

for i, row in enumerate(filters):
    add_table_data_row(table12, [(row[0], True), (row[1], True), row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_info_box(doc,
    'Performance: Database Indexes',
    [
        'All filter columns are individually indexed in the database:',
        '  • resources_status_idx          (status)',
        '  • resources_level_idx           (level)',
        '  • resources_resource_type_idx   (resource_type)',
        '  • resources_subject_idx         (subject)',
        '  • resources_year_idx            (year)',
        '  • resources_topic_idx           (topic)',
        '  • resources_language_idx        (language)',
        '  • resources_tags_gin_idx        (tags, GIN index for array search)',
        '  • resources_uploaded_by_idx     (uploaded_by)',
    ],
    bg_hex='F0FDF4',
    border_hex='86EFAC',
    title_color=GREEN,
)


# ── 8. Mobile & Performance Architecture ─────────────────────
doc.add_page_break()
add_heading(doc, '8. Mobile & Performance Architecture', level=1)
add_section_divider(doc)

add_body(doc,
    'The majority of Ubumenyi users will access the platform on Android smartphones over 3G or '
    'mobile data connections. Every design and technical decision is made with this constraint '
    'as the primary requirement, not an afterthought.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '8.1 Performance Strategies', level=2)
doc.add_paragraph()

perf_items = [
    ('Server-Side Rendering (SSR)', 'All public pages (resource library, detail pages, contributor profiles) are server-rendered — the browser receives complete HTML with no JavaScript required to see content. This is critical for low-powered devices.'),
    ('No Client-Side Filtering', 'All search and filter operations happen at the database level (PostgreSQL), not in the browser. This eliminates large data transfers.'),
    ('Minimal JavaScript Bundle', 'Tailwind CSS v4 with utility classes produces a minimal CSS bundle. No animation libraries or heavy UI frameworks.'),
    ('Image Optimization', 'Next.js built-in image optimization with WebP conversion and lazy loading for any images served.'),
    ('Supabase Edge Functions', 'Database queries execute close to the user via Supabase\'s edge infrastructure, reducing latency for users in Rwanda.'),
    ('CDN via Vercel', 'Static assets and pages are distributed via Vercel\'s global CDN — African edge nodes reduce time-to-first-byte.'),
    ('PDF Lazy Loading', 'The embedded PDF viewer on resource detail pages only loads when the user scrolls to it — initial page load is instant.'),
    ('Compressed Assets', 'All uploaded PDFs are stored as-is but served via CDN with compression headers. File size metadata is displayed to help users on metered connections.'),
]

table13 = doc.add_table(rows=1, cols=2)
table13.style = 'Table Grid'
table13.autofit = False
table13.columns[0].width = Inches(2.2)
table13.columns[1].width = Inches(4.2)
add_table_header_row(table13, ['Strategy', 'Implementation'])

for i, row in enumerate(perf_items):
    add_table_data_row(table13, [(row[0], True), row[1]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_heading(doc, '8.2 Mobile Design Principles', level=2)
doc.add_paragraph()

mobile_principles = [
    'All layouts are mobile-first using Tailwind CSS responsive breakpoints',
    'Navigation collapses to a hamburger menu with a slide-in drawer on mobile',
    'Filter sidebars stack vertically above results on screens below 1024px',
    'Resource cards use a single-column grid on mobile, 2-column on tablet, 3-column on desktop',
    'Admin panel includes a mobile-accessible top bar with a slide-out navigation drawer',
    'All touch targets meet minimum 44×44px size requirements for mobile usability',
    'Tables in the admin panel scroll horizontally on small screens (overflow-x: auto)',
    'File size is always displayed so users on metered connections can make informed choices',
]

for item in mobile_principles:
    add_bullet(doc, item)

doc.add_paragraph()


# ── 9. Design System ──────────────────────────────────────────
add_heading(doc, '9. Design System & Branding', level=1)
add_section_divider(doc)

add_body(doc,
    'Ubumenyi\'s visual identity is designed to feel trustworthy, institutional, and academic — '
    'not like a consumer startup or entertainment app. The design language borrows from digital '
    'library and government digital service patterns.',
    size=10.5
)
doc.add_paragraph()

colors_table = doc.add_table(rows=1, cols=3)
colors_table.style = 'Table Grid'
colors_table.autofit = False
colors_table.columns[0].width = Inches(1.6)
colors_table.columns[1].width = Inches(1.4)
colors_table.columns[2].width = Inches(3.4)
add_table_header_row(colors_table, ['Token', 'Hex Value', 'Usage'])

colors = [
    ('Brand Blue (Primary)', '#3457A6', 'Primary actions, links, headings, brand elements, active states'),
    ('Near-Black', '#0F172A', 'Primary heading text, high-emphasis content'),
    ('Slate 700', '#334155', 'Body text, card content'),
    ('Slate 500', '#64748B', 'Secondary text, labels, metadata'),
    ('Slate 400', '#94A3B8', 'Placeholder text, tertiary labels, timestamps'),
    ('Border Gray', '#E2E8F0', 'Card borders, table borders, dividers'),
    ('Background Gray', '#F8FAFC', 'Page backgrounds, table row alternates, input backgrounds'),
    ('Emerald (Success)', '#059669', 'Published status, approve actions, success states'),
    ('Amber (Warning)', '#D97706', 'Pending review status, moderation alerts'),
    ('Red (Error)', '#DC2626', 'Rejected status, delete actions, error states'),
    ('Indigo', '#5B21B6', 'Teacher notes resource type badge'),
    ('Teal', '#0A7068', 'Marking scheme resource type badge'),
]

for i, row in enumerate(colors):
    add_table_data_row(colors_table, [(row[0], True), row[1], row[2]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()

add_info_box(doc, 'Typography', [
    'Font: Geist Sans (via Next.js font optimization)',
    'Body text: 14px / 1.6 line-height',
    'Small labels / metadata: 10–12px',
    'Section headings: 18–24px, font-weight 600–700',
    'All text rendered at native system sharpness — no decorative fonts',
])


# ── 10. Scalability & Future Architecture ─────────────────────
add_heading(doc, '10. Scalability & Future Architecture', level=1)
add_section_divider(doc)

add_body(doc,
    'The platform is designed to scale from a small initial collection to a national-scale '
    'educational resource repository. Current architectural decisions support this trajectory.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '10.1 Current Capacity', level=2)
doc.add_paragraph()

capacity = [
    'PostgreSQL (Supabase) supports millions of rows — no practical resource count limit',
    'GIN indexes on the tags array column support efficient tag-based search at scale',
    'Supabase Storage supports up to 50MB per file with no total storage limit (paid tier)',
    'Vercel serverless deployment auto-scales to traffic — no manual capacity planning',
    'Row-Level Security removes performance-costly application-layer access checks',
]

for item in capacity:
    add_bullet(doc, item)

doc.add_paragraph()

add_heading(doc, '10.2 Planned Future Features', level=2)
doc.add_paragraph()

future = [
    ('User Accounts (Students)', 'Allow students to create accounts to save resources, track viewed materials, and receive topic recommendations'),
    ('Umukoro Integration', 'Connect the Ubumenyi resource library to the Umukoro interactive assessment platform for past paper practice'),
    ('Full-Text Search', 'PostgreSQL tsvector full-text search across resource titles, descriptions, and uploaded text content'),
    ('Multilingual Interface', 'Kinyarwanda and French UI translations for wider accessibility across Rwanda'),
    ('Offline Support (PWA)', 'Progressive Web App capabilities to allow downloaded resources to be accessed offline on Android'),
    ('Contributor Dashboard', 'Self-service portal for verified contributors to track their upload stats and manage submissions'),
    ('API Access', 'Public read API for REB systems and partner institutions to query the resource catalog programmatically'),
    ('Analytics Dashboard', 'Admin-facing analytics: most-downloaded resources, trending subjects, geographic access patterns'),
]

table14 = doc.add_table(rows=1, cols=2)
table14.style = 'Table Grid'
table14.autofit = False
table14.columns[0].width = Inches(2.0)
table14.columns[1].width = Inches(4.4)
add_table_header_row(table14, ['Feature', 'Description'])

for i, row in enumerate(future):
    add_table_data_row(table14, [(row[0], True), row[1]], i+1, alt=(i % 2 == 0))

doc.add_paragraph()


# ── 11. Component Architecture ────────────────────────────────
add_heading(doc, '11. Frontend Component Architecture', level=1)
add_section_divider(doc)

add_body(doc,
    'The application is structured as a Next.js 15 App Router project with a clear separation '
    'between public-facing pages, admin pages, shared components, server actions, and type definitions.',
    size=10.5
)
doc.add_paragraph()

add_heading(doc, '11.1 Directory Structure', level=2)
doc.add_paragraph()

structure = doc.add_paragraph()
structure.paragraph_format.left_indent = Inches(0.4)
r = structure.add_run(
    'src/\n'
    '  app/\n'
    '    page.tsx               — Landing page\n'
    '    resources/\n'
    '      page.tsx             — Resource library with search/filter\n'
    '      [id]/page.tsx        — Resource detail + PDF viewer\n'
    '    contributors/\n'
    '      page.tsx             — Public contributor directory\n'
    '      [id]/page.tsx        — Contributor profile page\n'
    '    admin/\n'
    '      layout.tsx           — Admin shell with sidebar + mobile top bar\n'
    '      page.tsx             — Admin dashboard\n'
    '      resources/page.tsx   — Resource management table\n'
    '      resources/new/       — Upload form\n'
    '      resources/[id]/edit/ — Edit form\n'
    '      contributors/page.tsx — Contributor moderation\n'
    '      team/page.tsx        — Admin team management\n'
    '    login/page.tsx         — Auth page\n'
    '    api/\n'
    '      confirm-admin/       — Admin confirmation API route\n'
    '  components/\n'
    '    landing/               — Landing page section components\n'
    '    admin/                 — Admin UI components (Sidebar, ResourceForm, etc.)\n'
    '    ui/                    — shadcn/ui primitive components\n'
    '  actions/\n'
    '    resources.ts           — Server actions: publish, reject, delete, contributor approval\n'
    '  lib/supabase/\n'
    '    client.ts              — Browser Supabase client\n'
    '    server.ts              — Server-side Supabase client (SSR cookies)\n'
    '    types.ts               — Auto-generated Supabase TypeScript types\n'
    '  types/\n'
    '    index.ts               — Application types, enums, label maps, constants\n'
    '  middleware.ts             — Route protection for /admin routes\n'
    'supabase/\n'
    '  migrations/              — SQL migration files\n'
    '  resources-table.sql      — Initial resources table\n'
)
r.font.name = 'Courier New'
r.font.size = Pt(8.5)
r.font.color.rgb = MID

doc.add_paragraph()

# ── 12. Summary ───────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '12. Summary', level=1)
add_section_divider(doc)

add_body(doc,
    'Ubumenyi is built with the architecture, security, and quality standards appropriate for a '
    'national educational infrastructure platform. Every technical decision has been made with '
    'three constraints in mind: trustworthiness for institutional adoption, accessibility for '
    'students on mobile devices in low-bandwidth environments, and scalability to support '
    'Rwanda\'s full educational resource ecosystem.',
    size=10.5
)

doc.add_paragraph()

summary_points = [
    'Database-enforced security via PostgreSQL Row-Level Security on every table',
    'Moderated contributor system — no unverified users can publish content',
    'Multi-stage resource workflow — every upload reviewed before publication',
    'Mobile-first, low-bandwidth-optimized server-rendered frontend',
    'Comprehensive classification schema (8 dimensions) enabling precise search and filtering',
    'Full audit trail: download logs, moderation timestamps, moderator attribution',
    'Scalable architecture with performance indexes — designed for 10,000+ resources',
    'Curriculum-aligned taxonomy organized to Rwanda\'s REB educational levels',
    'Free to access — no paywalls, no mandatory registration for resource browsing',
]

for point in summary_points:
    add_bullet(doc, point)

doc.add_paragraph()
doc.add_paragraph()

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_contact.add_run('For technical inquiries or partnership discussions:')
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = MUTED

p_email = doc.add_paragraph()
p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_email.add_run('hello@ubumenyi.rw   |   ubumenyi.rw   |   Kigali, Rwanda')
r2.font.size = Pt(11)
r2.font.bold = True
r2.font.name = 'Calibri'
r2.font.color.rgb = BLUE

# ── Save ──────────────────────────────────────────────────────
output_path = '/sessions/loving-fervent-lovelace/mnt/outputs/Ubumenyi_Technical_Architecture_v2.docx'
doc.save(output_path)
print(f"Document saved: {output_path}")
